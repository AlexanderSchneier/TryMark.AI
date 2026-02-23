from document import create_document
from knowledge.retrieval import retrieve_relevant_chunks_for_section
from generation.payload_builder import build_generation_payload
from generation.generate import generate_text
from docx import Document
from io import BytesIO
from dotenv import load_dotenv

load_dotenv()

from main import SECTIONS


# -------------------------------------------------------------------
# Mandatory clause templates (verbatim injection)
# -------------------------------------------------------------------
MANDATORY_CLAUSES = {
    # Foundational sweepstakes classification clauses
    "HC-002": "Winner will be selected in a random drawing from all eligible entries received during the Promotion Period.",
    "HC-003": "All eligible entries will have an equal chance of winning, regardless of method of entry.",

    # Triggered bonding/registration clause
    "HC-013": "If required by applicable law, Sponsor will register and bond this Sweepstakes in Florida and/or New York, as applicable, prior to awarding prizes.",
}


def _select_required_clauses_for_section(constraint_output: dict, section_category: str) -> list[dict]:
    """
    Decide which mandatory clauses must appear in *this* section.

    We intentionally route:
      - HC-002/HC-003 -> sweepstakes_classification section
      - HC-013 -> bonding_registration (your General Conditions section)
    """
    required: list[dict] = []

    # include both foundational + triggered rules
    for group_name in ("foundational", "triggered"):
        for r in constraint_output.get(group_name, []):
            cid = r.get("id")
            if not cid:
                continue
            if cid not in MANDATORY_CLAUSES:
                continue

            if cid in ("HC-002", "HC-003") and section_category == "sweepstakes_classification":
                required.append({"id": cid, "text": MANDATORY_CLAUSES[cid]})

            if cid == "HC-013" and section_category == "bonding_registration":
                required.append({"id": cid, "text": MANDATORY_CLAUSES[cid]})

    return required


def _missing_required_clauses(section_text: str, required_clauses: list[dict]) -> list[dict]:
    if not required_clauses:
        return []
    text_lower = (section_text or "").lower()
    missing = []
    for c in required_clauses:
        clause_text = (c.get("text") or "").strip()
        if clause_text and clause_text.lower() not in text_lower:
            missing.append(c)
    return missing


def _looks_truncated(section_text: str) -> bool:
    """
    Cheap truncation heuristic to catch outputs that end mid-sentence.
    """
    t = (section_text or "").strip()
    if not t:
        return True
    tail = t[-40:].lower()
    return (
        t.endswith(",")
        or t.endswith(":")
        or "at its sole discretion," in tail
        or "at its sole discretion" == t.lower().strip()
    )


def generate_official_rules(form_data: dict):

    # Build document using provided data instead of CLI prompts
    doc = create_document(from_api_data=form_data)
    doc.load_hard_constraints("hard_constraints.json")
    doc.apply_hard_constraints()
    doc.validate()

    promotion_context = {
        "name": doc._name,
        "states": doc._residence,
        "min_age": doc._minAge,
        "start_time": doc._startTime,
        "end_time": doc._endTime,
        "winner_selection_time": doc._winnerTime,
        "winner_response_deadline": doc._winnerResponseTime,
        "primary_prize_type": doc._prizes.value,
        "prizes": [
            (
                {"type": p.prize_type.value, "amount": p.amount}
                if p.prize_type.value == "cash"
                else {"type": p.prize_type.value, "description": p.description}
            )
            for p in doc._prizeLevels.values()
        ],
        "total_prize_value": doc._total_prize_value(),

        # 🔥 NEW STRUCTURED ENTRY METHOD
        "entry_method": {
            "channel": doc._entryChannel,
            "url": doc._entryUrl,
            "required_fields": doc._entryFields
        },

        "in_store_entry": doc._inPersonEntry
    }

    compliance_requirements = doc._constraint_output
    generated_sections: dict[str, str] = {}

    for section in SECTIONS:

        section_category = section["category"]
        section_title = section["title"]

        # SECTION-AWARE RETRIEVAL
        relevant_snippets = retrieve_relevant_chunks_for_section(
            compliance_requirements,
            section_category=section_category,
            section_title=section_title,
            top_k=6,
            always_include_baseline=True,
            min_score=15
        )

        # ✅ Mandatory clauses for this section
        required_clauses = _select_required_clauses_for_section(
            compliance_requirements,
            section_category=section_category
        )

        # Build Payload (now includes required_clauses)
        payload = build_generation_payload(
            promotion_context=promotion_context,
            compliance_requirements=compliance_requirements,
            historical_snippets=relevant_snippets,
            section_name=section_title,
            section_category=section_category,
            required_clauses=required_clauses
        )


        # ---- Generate with 1 retry if enforcement fails ----
        section_text = generate_text(payload)

        missing = _missing_required_clauses(section_text, required_clauses)
        truncated = _looks_truncated(section_text)

        # Retry once if needed
        if missing or truncated:
            extra = "\n\nCORRECTION REQUIRED:\n"

            if missing:
                extra += "You omitted mandatory clause(s). You MUST include each clause EXACTLY as written:\n"
                for c in missing:
                    extra += f"- [{c['id']}] {c['text']}\n"

            if truncated:
                extra += "Your section appears cut off. You MUST provide a complete section ending with a full sentence.\n"

            payload_retry = {"prompt": payload["prompt"] + extra}
            section_text = generate_text(payload_retry)

        # 🔒 FAIL-CLOSED ENFORCEMENT (Deterministic Append)
        final_missing = _missing_required_clauses(section_text, required_clauses)

        if final_missing:
            for c in final_missing:
                # Append clause directly if model failed to include it
                section_text += f"\n\n{c['text']}\n"
        generated_sections[section["id"]] = section_text





    # Build docx in memory
    document = Document()
    document.add_heading("OFFICIAL SWEEPSTAKES RULES", level=1)

    for section in SECTIONS:
        document.add_heading(section["title"], level=2)
        content = generated_sections.get(section["id"], "")

        for line in content.split("\n"):
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer