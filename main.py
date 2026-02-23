from dotenv import load_dotenv
load_dotenv()

from document import create_document
from knowledge.retrieval import retrieve_relevant_chunks_for_section
from generation.payload_builder import build_generation_payload
from generation.generate import generate_text
import json
from docx import Document


# Canonical Official Rules structure
SECTIONS = [
    {
        "id": "classification",
        "title": "Agreement to Official Rules",
        "category": "sweepstakes_classification"
    },
    {
        "id": "eligibility",
        "title": "Eligibility",
        "category": "eligibility"
    },
    {
        "id": "entry_method",
        "title": "How to Enter",
        "category": "entry_method"
    },
    {
        "id": "prizes",
        "title": "Prize(s)",
        "category": "prizes"
    },
    {
        "id": "winner_clearance",
        "title": "Requirements of Potential Winners",
        "category": "winner_clearance"
    },
    {
        "id": "general_conditions",
        "title": "General Conditions",
        "category": "bonding_registration"
    }
]


def main():

    # 1️⃣ Collect promotion facts
    doc = create_document()
    doc.load_hard_constraints("hard_constraints.json")
    doc.apply_hard_constraints()
    doc.validate()

    promotion_context = {
        "name": doc._name,
        "states": doc._residence,
        "min_age": doc._minAge,

        "start_time": doc._startTime,
        "end_time": doc._endTime,
        "winner_selection_time": doc._winnerTime,
        "winner_response_deadline": doc._winnerResponseTime,

        "prizes": [
            {"type": p.prize_type.value, "amount": p.amount}
            for p in doc._prizeLevels.values()
        ],

        "total_prize_value": doc._total_prize_value(),

        # You’ll replace this later with real logic
        "entry_method": "unspecified",
        "in_store_entry": doc._inPersonEntry
    }

    compliance_requirements = doc._constraint_output

    print("\n=== CONSTRAINT OUTPUT ===\n")
    print(json.dumps(compliance_requirements, indent=2))


    # 2️⃣ Generate document section-by-section (SECTION-AWARE RAG)
    print("\n=== GENERATING DOCUMENT SECTIONS ===\n")

    generated_sections = {}

    for section in SECTIONS:

        category = section["category"]
        title = section["title"]

        # 🔥 SECTION-AWARE RETRIEVAL
        relevant_snippets = retrieve_relevant_chunks_for_section(
            compliance_requirements,
            section_category=category,
            section_title=title,
            top_k=6,                    # Tune between 4–8
            always_include_baseline=True,
            min_score=15
        )

        # ---- Debug Output (Very Important For Tuning) ----
        print(f"\n=== RAG FOR: {title} ({category}) ===")

        if not relevant_snippets:
            print("  ⚠️ No snippets retrieved.")
        else:
            for i, c in enumerate(relevant_snippets, 1):
                print(
                    f"  [{i}] {c.get('id')} | "
                    f"section={c.get('section')} | "
                    f"score={c.get('_score')} | "
                    f"hard={c.get('hard_constraint')}"
                )
                preview = c.get("text", "")[:160].replace("\n", " ")
                print(f"      preview: {preview}")

        # ---- Build Payload ----
        payload = build_generation_payload(
            promotion_context,
            compliance_requirements,
            relevant_snippets,
            section_name=title,
            section_category=category
        )

        print(f"→ Generating section: {title}")

        section_text = generate_text(payload)
        generated_sections[section["id"]] = section_text

    print("\n\n=== FINAL GENERATED DOCUMENT ===\n")

    document = Document()

    # Optional: Title Page Header
    document.add_heading("OFFICIAL SWEEPSTAKES RULES", level=1)

    for section in SECTIONS:
        title = section["title"]
        content = generated_sections.get(section["id"], "")

        # Add section title as proper Word heading
        document.add_heading(title, level=2)

        # Preserve line breaks from model output
        for line in content.split("\n"):
            document.add_paragraph(line)

    # Dynamic filename using sweepstakes name
    safe_name = promotion_context["name"].replace(" ", "_")
    filename = f"{safe_name}_Official_Rules.docx"

    document.save(filename)

    print(f"✅ Document saved as {filename}")

if __name__ == "__main__":
    main()
